from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "melee-design-packet-v0.5.md"
TEMPLATE = ROOT / "Atra_Melee_Design_Packet_v0.4.docx"
OUTPUT = ROOT / "Atra_Melee_Design_Packet_v0.5.docx"


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def strip_markdown(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def set_headers(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    p = header.paragraphs[0]
    p.clear()
    p.text = "ATRA RPG  |  MELEE DESIGN PACKET v0.5"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Light Shading Accent 1"
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col_index in range(width):
            value = values[col_index] if col_index < len(values) else ""
            cell = row.cells[col_index]
            cell.text = strip_markdown(value.strip())
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True
        if row_index == 0:
            set_cell_repeat_header(row)
    document.add_paragraph()


def build() -> None:
    shutil.copyfile(TEMPLATE, OUTPUT)
    document = Document(OUTPUT)
    clear_body(document)
    configure_styles(document)
    set_headers(document)

    props = document.core_properties
    props.title = "Atra Melee Design Packet v0.5"
    props.subject = "Current Governing-Provisional Vertical Slice"
    props.author = "Atra RPG Project"
    props.comments = "Generated from docs/melee-design-packet-v0.5.md using the preserved v0.4 DOCX style authority."

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    table_buffer: list[list[str]] = []
    first_h1 = True
    cover_complete = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                index += 1
                continue
            table_buffer.append(cells)
            index += 1
            continue
        flush_table()

        if not line:
            index += 1
            continue
        if line.startswith("# "):
            title = line[2:]
            if first_h1:
                p = document.add_paragraph(style="Title")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_inline(p, title)
                first_h1 = False
            else:
                document.add_paragraph(title, style="Heading 1")
            index += 1
            continue
        if line.startswith("## "):
            text = line[3:]
            if not cover_complete and text == "Current Governing-Provisional Vertical Slice":
                p = document.add_paragraph(style="Subtitle")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_inline(p, text)
            else:
                if not cover_complete:
                    document.add_page_break()
                    cover_complete = True
                else:
                    document.add_page_break()
                document.add_paragraph(text, style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 2")
            index += 1
            continue
        if line.startswith("#### "):
            document.add_paragraph(line[5:], style="Heading 3")
            index += 1
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            p = document.add_paragraph(style="List Number")
            add_inline(p, numbered.group(2))
            index += 1
            continue
        if re.fullmatch(r"-+", line):
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            nxt = lines[index].rstrip()
            if not nxt or nxt.startswith(("#", "- ", "|")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = document.add_paragraph()
        add_inline(p, " ".join(paragraph_lines).replace("  ", " "))

    flush_table()
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
